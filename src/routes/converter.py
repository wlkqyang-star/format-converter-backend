import os
import json
import csv
import io
import tempfile
import markdown2
from docx import Document
from docx.shared import Inches
import re
from PIL import Image
import pillow_heif
# from moviepy.editor import VideoFileClip # Temporarily removed due to import issues
from flask import Blueprint, request, jsonify, send_file
from werkzeug.utils import secure_filename
import uuid
import pytesseract
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader

# 注册HEIF插件
pillow_heif.register_heif_opener()

converter_bp = Blueprint("converter", __name__)

# 允许的文件扩展名
ALLOWED_EXTENSIONS = {
    "json": ["json"],
    "csv": ["csv"],
    "markdown": ["md", "markdown"],
    "docx": ["docx"],
    "image": ["jpg", "jpeg", "png", "heic", "webp", "avif"],
    "video": ["mp4", "avi", "mov", "mkv"],
    "pdf": ["pdf"],
}

def allowed_file(filename, file_type):
    """检查文件扩展名是否被允许"""
    return "." in filename and \
           filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS.get(file_type, [])

def json_to_csv(json_data):
    """将JSON数据转换为CSV格式"""
    try:
        if isinstance(json_data, str):
            data = json.loads(json_data)
        else:
            data = json_data
        
        if isinstance(data, dict):
            data = [data]
        
        if not isinstance(data, list):
            raise ValueError("JSON数据必须是对象数组或单个对象")
        
        if not data:
            return ""
        
        fieldnames = set()
        for item in data:
            if isinstance(item, dict):
                fieldnames.update(item.keys())
        
        fieldnames = sorted(list(fieldnames))
        
        output = io.StringIO(newline="")
        writer = csv.DictWriter(output, fieldnames=fieldnames)
        output.write("\ufeff") # UTF-8 BOM for Excel
        writer.writeheader()
        
        for item in data:
            if isinstance(item, dict):
                row = {field: item.get(field, "") for field in fieldnames}
                writer.writerow(row)
        
        return output.getvalue()
    
    except Exception as e:
        raise ValueError(f"JSON转CSV失败: {str(e)}")

def csv_to_json(csv_data):
    """将CSV数据转换为JSON格式"""
    try:
        csv_file = io.StringIO(csv_data)
        reader = csv.DictReader(csv_file)
        result = [row for row in reader]
        return json.dumps(result, ensure_ascii=False, indent=2)
    except Exception as e:
        raise ValueError(f"CSV转JSON失败: {str(e)}")

def markdown_to_docx(markdown_content):
    """将Markdown内容转换为DOCX格式"""
    try:
        html = markdown2.markdown(markdown_content, extras=["fenced-code-blocks", "tables"])
        doc = Document()
        # This is a simplified conversion, can be extended
        lines = html.split("\n")
        for line in lines:
            clean_line = re.sub(r"<[^>]+>", "", line.strip())
            if clean_line:
                doc.add_paragraph(clean_line)
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        temp_file.close()
        return temp_file.name
    except Exception as e:
        raise ValueError(f"Markdown转DOCX失败: {str(e)}")

def docx_to_markdown(docx_file_path):
    """将DOCX文件转换为Markdown格式"""
    try:
        doc = Document(docx_file_path)
        markdown_content = [p.text for p in doc.paragraphs]
        return "\n\n".join(markdown_content)
    except Exception as e:
        raise ValueError(f"DOCX转Markdown失败: {str(e)}")

def convert_image_format(input_file_path, target_format):
    """转换图片格式"""
    try:
        with Image.open(input_file_path) as img:
            if target_format.upper() in ["JPEG", "JPG"] and img.mode in ["RGBA", "LA"]:
                background = Image.new("RGB", img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[-1])
                img = background
            
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f".{target_format.lower()}")
            temp_file.close()
            img.save(temp_file.name, format=target_format.upper())
            return temp_file.name
    except Exception as e:
        raise ValueError(f"图片格式转换失败: {str(e)}")

# def video_to_gif(video_file_path):
#     """将视频文件转换为GIF格式"""
#     try:
#         clip = VideoFileClip(video_file_path)
#         temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".gif")
#         temp_file.close()
#         clip.write_gif(temp_file.name, fps=10) # fps can be adjusted
#         clip.close()
#         return temp_file.name
#     except Exception as e:
#         raise ValueError(f"视频转GIF失败: {str(e)}")

def image_to_text_ocr(image_path, lang='chi_sim'):
    """使用OCR从图片中提取文本"""
    try:
        text = pytesseract.image_to_string(Image.open(image_path), lang=lang)
        return text
    except Exception as e:
        raise ValueError(f"图片OCR识别失败: {str(e)}")

def image_to_searchable_pdf_ocr(image_path, lang='chi_sim'):
    """将图片转换为可搜索的PDF"""
    try:
        pdf_file = pytesseract.image_to_pdf_or_hocr(Image.open(image_path), lang=lang, extension='pdf')
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        temp_file.write(pdf_file)
        temp_file.close()
        return temp_file.name
    except Exception as e:
        raise ValueError(f"图片转可搜索PDF失败: {str(e)}")

@converter_bp.route("/convert/<conversion_type>", methods=["POST"])
def handle_conversion(conversion_type):
    if "file" not in request.files:
        return jsonify({"error": "没有上传文件"}), 400
    
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "没有选择文件"}), 400

    try:
        filename = secure_filename(file.filename)
        file_ext = filename.rsplit(".", 1)[1].lower() if "." in filename else ""
        unique_id = str(uuid.uuid4())[:8]

        # Save uploaded file to a temporary location
        temp_input = tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_ext}")
        file.save(temp_input.name)
        temp_input.close()

        output_file_path = None
        output_filename = None
        mimetype = None

        if conversion_type == "json-csv":
            if file_ext == "json":
                content = open(temp_input.name, "r", encoding="utf-8").read()
                csv_content = json_to_csv(content)
                output_file_path = tempfile.NamedTemporaryFile(mode="w", delete=False, suffix=".csv", encoding="utf-8").name
                with open(output_file_path, "w", encoding="utf-8") as f:
                    f.write(csv_content)
                output_filename = f"converted_{unique_id}.csv"
                mimetype = "text/csv"
            elif file_ext == "csv":
                content = open(temp_input.name, "r", encoding="utf-8").read()
                json_content = csv_to_json(content)
                output_file_path = tempfile.NamedTemporaryFile(mode="w", delete=False, suffix=".json", encoding="utf-8").name
                with open(output_file_path, "w", encoding="utf-8") as f:
                    f.write(json_content)
                output_filename = f"converted_{unique_id}.json"
                mimetype = "application/json"

        elif conversion_type == "markdown-docx":
            if file_ext in ["md", "markdown"]:
                content = open(temp_input.name, "r", encoding="utf-8").read()
                output_file_path = markdown_to_docx(content)
                output_filename = f"converted_{unique_id}.docx"
                mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif file_ext == "docx":
                output_content = docx_to_markdown(temp_input.name)
                output_file_path = tempfile.NamedTemporaryFile(mode="w", delete=False, suffix=".md", encoding="utf-8").name
                with open(output_file_path, "w", encoding="utf-8") as f:
                    f.write(output_content)
                output_filename = f"converted_{unique_id}.md"
                mimetype = "text/markdown"

        elif conversion_type == "image-format":
            target_format = request.form.get("target_format", "png").lower()
            output_file_path = convert_image_format(temp_input.name, target_format)
            output_filename = f"converted_{unique_id}.{target_format}"
            mimetype = f"image/{target_format}"

# elif conversion_type == "video-gif": # Temporarily removed
#     output_file_path = video_to_gif(temp_input.name)
#     output_filename = f"converted_{unique_id}.gif"
#     mimetype = "image/gif"

        elif conversion_type == "image-to-text-ocr":
            if file_ext in ALLOWED_EXTENSIONS["image"]:
                text_content = image_to_text_ocr(temp_input.name)
                output_file_path = tempfile.NamedTemporaryFile(mode="w", delete=False, suffix=".txt", encoding="utf-8").name
                with open(output_file_path, "w", encoding="utf-8") as f:
                    f.write(text_content)
                output_filename = f"converted_{unique_id}.txt"
                mimetype = "text/plain"
            else:
                return jsonify({"error": "不支持的图片格式进行OCR识别"}), 400

        elif conversion_type == "image-to-searchable-pdf-ocr":
            if file_ext in ALLOWED_EXTENSIONS["image"]:
                output_file_path = image_to_searchable_pdf_ocr(temp_input.name)
                output_filename = f"converted_{unique_id}.pdf"
                mimetype = "application/pdf"
            else:
                return jsonify({"error": "不支持的图片格式转换为可搜索PDF"}), 400

        if output_file_path and output_filename and mimetype:
            response = send_file(output_file_path, as_attachment=True, download_name=output_filename, mimetype=mimetype)
            os.unlink(output_file_path)
            return response
        else:
            return jsonify({"error": "不支持的转换类型或文件格式"}), 400

    except Exception as e:
        return jsonify({"error": f"转换失败: {str(e)}"}), 500
    finally:
        if "temp_input" in locals() and os.path.exists(temp_input.name):
            os.unlink(temp_input.name)




